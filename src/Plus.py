'''
Created on May 28, 2016

@author: wange
'''
from __future__ import print_function

import numpy as np
from docx import Document 
from docx.shared import Pt
from docx.text import parfmt
from random import shuffle
import random
from sys import argv

from reportlab.pdfgen import canvas

point = 1
inch = 72

def plus(mini,limited, Neqns):
    eqn_ele=[]
    for i in range(0,Neqns):
        a=np.random.randint(mini,limited)
        b=np.random.randint(mini,limited)
        #a=np.random.randint(1,10)
        #b=np.random.randint(1,10)
        i+=1
        eqn_ele.append('%d+%d=        '%(a,b))
        #eqn_ele.append('\n')
    #print(eqn_ele)
    return eqn_ele

def minus(mini,limited, Neqns):
    eqn_ele=[]
    for i in range(0,Neqns):
        a=np.random.randint(mini,limited)
        b=np.random.randint(5,a)
        i+=1
        eqn_ele.append('%d-%d=         '%(a,b))
        #eqn_ele.append('\n')
    #print(eqn_ele)
    return eqn_ele

def times(mini,limited, Neqns):
    eqn_ele=[]
    for i in range(0,Neqns):
        a=np.random.randint(mini,limited)
        b=np.random.randint(mini,limited)
        i+=1
        eqn_ele.append('%d'%(a)+u"\u00D7"+'%d=            '%(b))
    return eqn_ele
        
def divide(mini,limited, Neqns):
    eqn_ele=[]
    for i in range(0,Neqns):
        
        b=np.random.randint(mini,limited)
        c=np.random.randint(2,mini)
        a=b*c+np.random.randint(0,c)
        i+=1
        eqn_ele.append('%d'%(a)+u"\u00F7"+'%d=         '%(c))
        #eqn_ele.append('\n')
    #print(eqn_ele)
    return eqn_ele

def equation(limited,Neqns):
    eqn_ele=[]
    signcal=["+","-","x",":"]
    unknownum=["X","Y","Z","q","p"]
 
    for i in range(0,Neqns):
        unknowlet=random.choice(unknownum)
        signeq=random.choice(signcal)
        #print(signeq)
        a=np.random.randint(2,limited)
        if signeq=="+":
            if np.random.rand()>0.5:
                eqn_ele.append('%d '%(a)+signeq+unknowlet+" =%d        "%(a+np.random.randint(2,limited)))
            else:
                eqn_ele.append(unknowlet+signeq+'%d '%(a)+" =%d        "%(a+np.random.randint(2,limited)))
                
        if signeq==":":
            if np.random.rand()>0.5:
                eqn_ele.append(unknowlet+u"\u00F7"+'%d '%(a)+" = %d        "%(np.random.randint(2,limited)))            
            else:
                eqn_ele.append('%d '%(a+np.random.randint(2,limited))+u"\u00F7"+unknowlet+" =%d        "%(a))               
                
        if signeq=="-":
            if np.random.rand()>0.5:
                eqn_ele.append('%d '%(a)+signeq+unknowlet+" =%d        "%(np.random.randint(2,a)))
            else:
                eqn_ele.append(unknowlet+signeq+'%d '%(a)+" =%d        "%(a+np.random.randint(2,limited)))
                
        if signeq=="x":
                eqn_ele.append('%d '%(a)+u"\u00D7" +unknowlet+" =%d        "%(a*np.random.randint(2,limited)))

    return eqn_ele

def mixcal(mini,limited,Neqns):
    eqn_ele=[]
    signcal=["+","-",u"\u00D7",u"\u00F7"]
    unknownum=["X","Y","Z","a","p"]
    
 
    brac=["(",")"]
    for i in range(Neqns):
        unknowlet=random.choice(unknownum)
        sig_num=np.random.randint(2,5)
        sign_cho=random.sample(set(signcal),sig_num)
        sign_cho.append("=")
        ele_num=np.random.randint(2,100,size=len(sign_cho))
        ele_num=list(map(str,ele_num))
        ele_num.insert(np.random.randint(len(sign_cho)), unknowlet)
 

        for k in range(len(ele_num)-1):
            k+=1
            ele_num[(2*k-1):(2*k-1)]=sign_cho[k-1]
        #ele_num.append("\n")
        indi=[n for n,j in enumerate(ele_num) if (j=="-" or j=="+")]
        
        if np.random.rand()>0.5:
            #print(ele_num,len(indi),indi)
            if len(indi)==1:
                ele_num[indi[0]-1:indi[0]-1]=brac[0]
                ele_num[indi[0]+3:indi[0]+3]=brac[1]
            elif len(indi)==2:
                brac_pos=np.random.randint(0,1)
                ele_num[indi[brac_pos]-1:indi[brac_pos]-1]=brac[0]
                ele_num[indi[brac_pos]+3:indi[brac_pos]+3]=brac[1]
            #print(indi,ele_num)   
        ele_num=' '.join(ele_num)
        eqn_ele.append(ele_num)
    print(eqn_ele)
    streqn='\t\n'.join(eqn_ele)
    print(streqn)
    return streqn
        
        
        
        
       
        #ele_num=ele_num+sign_cho
        #ele_num.insert(np.random.randint(len(sign_cho)), unknowlet)
        

def comb(mini,limited,Nplus,Nminus,Ntimes,Ndivide,Neqns):
    print(Nplus,Nminus,Ntimes,Ndivide,Neqns)
    comb=plus(mini,limited*8,Nplus)
    comb.extend(minus(mini,limited*8,Nminus))
    comb.extend(times(mini,limited,Ntimes))
    comb.extend(divide(mini,limited,Ndivide))
    comb.extend(equation(limited, Neqns))
    shuffle(comb)
    #print(comb)
    strcomb='\t  \t'.join(str(x)+' ' for x in comb)
    #print(strcomb)
    return strcomb


def genedoc(days,minical,limitedcal,pluscal,minuscal,timescal,dividecal,eqncal):
    d=Document()
    
    #font.font.name='Calibri'
    for day in range(1,days):
        calc_list=comb(minical,limitedcal,pluscal,minuscal,timescal,dividecal,eqncal)
        d.add_heading('The %d day; Using__________mins'%(day))
        d.add_paragraph(calc_list)
        d.add_page_break()
        day+=1
    style=d.styles['Normal']
    style.font.size=Pt(24)    
    d.save('../Practise_mix.docx')



def genedoc_eqn(days,minical,limitedcal,Neqns):
    d=Document()
    #font.font.name='Calibri'
    for day in range(1,days):
        calc_list=mixcal(minical, limitedcal, Neqns)
        #print(calc_list)
        d.add_heading('The %d day; Using__________mins'%(day))
        d.add_paragraph(calc_list)
        # d.add_page_break()
        day+=1
    style=d.styles['Normal']
    style.font.size=Pt(24)    
    d.save('../Practise_eqn.docx')


if __name__ == '__main__':
    test=0
    if(test==1):
        totalcal=5
        minical=3
        limitedcal=100
        days=12
        genedoc_eqn(days, minical, limitedcal, totalcal)
    elif(test==0):
        totalcal=34
        pluscal=9
        minuscal=9
        timescal=8
        eqncal=0
        dividecal=totalcal-pluscal-minuscal-timescal-eqncal
        minical=12
        limitedcal=100
        days=2
        #dailycal=comb(minical,limitedcal,pluscal,minuscal,timescal)
        #print(dailycal)
        #genepdf("day1.pdf", dailycal)
        print("gene=",pluscal,minuscal,timescal,dividecal,eqncal)
        genedoc(days,minical,limitedcal,pluscal,minuscal,timescal,dividecal,eqncal)
 
